import os
import zipfile
import re
import argparse
from docx import Document

def docx_to_html(docx_file, output_html):
    doc = Document(docx_file)
    chapters = []

    # First pass: collect chapter titles for TOC
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading 1"):
            title = para.text.strip()
            chapters.append(title)
    # Second pass: build per-chapter HTML fragments

    chapter_idx = 0
    current_fragment = None
    current_title = None
    chapter_files = []

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name

        # Chapter start
        if style.startswith("Heading 1"):
            if current_fragment is not None:
                chapter_files.append((chapter_idx, current_title, current_fragment))
            print("Chapter_idx: {}".format(chapter_idx))
            print(text)
            chapter_idx += 1            

    print(chapter_files)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert a .docx file to HTML")
    parser.add_argument("input_docx", nargs="?", default="HaraldHardrada.docx", help="Input .docx file")
    parser.add_argument("output_html", nargs="?", default="harald.html", help="Output HTML file")
    args = parser.parse_args()

    docx_to_html(args.input_docx, args.output_html)

    print(f"Created {args.output_html}")

import os
import zipfile
import re
import argparse
from docx import Document


def docx_to_html(docx_file, output_html):
    doc = Document(docx_file)
    chapters = []

    # First pass: collect chapter titles for TOC
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading 1") or para.text.strip().startswith("## "):
            title = para.text.strip().replace("## ", "")
            chapters.append(title)

    # Prepare output directory and base name
    out_dir = os.path.dirname(output_html) or "."
    os.makedirs(out_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(output_html))[0]

    # Second pass: build per-chapter HTML fragments
    chapter_idx = 0
    i = 0
    current_fragment = None
    current_title = None
    current_section_title = None
    chapter_files = []

    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        style = para.style.name

        # Chapter start
        if style.startswith("Heading 1") or text.startswith("## "):
            if current_fragment is not None:
                chapter_files.append((chapter_idx-1, current_title, current_fragment))
            chapter_idx += 1
            current_title = text.replace("## ", "")
            current_fragment = []
            current_section_title = None
            current_fragment.append(f"<section class='chapter' id='chapter-{chapter_idx-1}'>")
            current_fragment.append(f"<figure><img src=images/chapter_{chapter_idx-1}.png><figcaption>Chapter {chapter_idx-1} Illustration</figcaption></figure>")
            current_fragment.append(f"<div class='chapter-title'>{current_title}</div>")
            # Gloss: next paragraph in italics or parentheses
            if i+1 < len(doc.paragraphs):
                gloss_para = doc.paragraphs[i+1]
                gloss_text = gloss_para.text.strip()
                if gloss_para.style.name == "Intense Emphasis" or (gloss_text.startswith("(") and gloss_text.endswith(")")):
                    gloss = gloss_text.lstrip("(").rstrip(")")
                    current_fragment.append(f"<div class='chapter-gloss'>{gloss}</div>")
                    i += 1
        # Subheading
        elif style.startswith("Heading 2") or text.startswith("### "):
            subtitle = text.replace("### ", "")
            if current_fragment is None:
                current_fragment = []
            current_fragment.append(f"<h2>{subtitle}</h2>")
            current_section_title = subtitle    
        # Quote
        elif style == "Quote":
            if current_fragment is None:
                current_fragment = []
            current_fragment.append(f"<div class='poetry'>{text}</div>")
        # Horizontal rule
        elif text == "---":
            if current_fragment is None:
                current_fragment = []
            current_fragment.append("<hr>")
        # List
        elif style == "List Paragraph":
            if current_fragment is None:
                current_fragment = []
            current_fragment.append("<ul>")
            while i < len(doc.paragraphs) and doc.paragraphs[i].style.name == "List Paragraph":
                item = doc.paragraphs[i].text.strip()
                current_fragment.append(f"<li>{item}</li>")
                i += 1
            current_fragment.append("</ul>")
            i -= 1
        # Normal paragraph
        elif text:
            if current_fragment is None:
                current_fragment = []
            run_html = ""
            for run in para.runs:
                run_text = run.text
                if run.bold and run.italic:
                    run_html += f"<strong><em>{run_text}</em></strong>"
                elif run.bold:
                    run_html += f"<strong>{run_text}</strong>"
                elif run.italic:
                    run_html += f"<em>{run_text}</em>"
                else:
                    run_html += run_text
            if (current_section_title):
                current_fragment.append(f"<p class='section-{current_section_title}'>{run_html}</p>")
            else:
                current_fragment.append(f"<p>{run_html}</p>")
        i += 1

    if current_fragment is not None:
        chapter_files.append((chapter_idx, current_title, current_fragment))

    # Write chapter files and TOC
    toc_items = []
    for idx, title, fragment in chapter_files:
        chapter_filename = f"{base_name}_chapter_{idx}.html"
        chapter_path = os.path.join(out_dir, chapter_filename)
        full_html = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "<meta charset='UTF-8'>",
            f"<title>{title}</title>",
            "<link rel='stylesheet' href='css/harald.css'>",
            "</head>",
            "<body>",
            "<h1 class='book-title'>Harald Hardrada, The Last Viking</h1>",
            "<div class='container'>",
            f"<nav><a href='{os.path.basename(output_html)}'>Table of Contents</a></nav>",
        ]
        if idx > 1:
            prev_filename = f"{base_name}_chapter_{idx-1}.html"
            full_html.append(f"<nav><a href='{prev_filename}'>Previous Chapter</a></nav>")
        if idx < len(chapter_files)-1:
            next_filename = f"{base_name}_chapter_{idx+1}.html"
            full_html.append(f"<nav><a href='{next_filename}'>Next Chapter</a></nav>")
        #add forward/backward navigation
        full_html.extend(fragment)
        full_html.extend(["</div>", "</body>", "</html>"])
        with open(chapter_path, "w", encoding="utf-8") as f:
            f.write("\n".join(full_html))
        toc_items.append((idx, title, chapter_filename))

    # Create TOC file linking to chapters
    toc_html = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "<meta charset='UTF-8'>",
        "<title>Harald Hardrada, The Last Viking</title>",
        "<link rel='stylesheet' href='css/harald.css'>",
        "</head>",
        "<body>",
        "<div class='container'>",
        "<h1 class='book-title'>Harald Hardrada, The Last Viking</h1>",
        "<nav class='toc'><h2>Table of Contents</h2><ul>"
    ]
    for idx, title, filename in toc_items:
        toc_html.append(f"<li><a href='{filename}'>{title}</a></li>")
    toc_html.extend(["</ul></nav>", "</div>", "</body>", "</html>"])

    with open(output_html, "w", encoding="utf-8") as f:
        f.write("\n".join(toc_html))



if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert a .docx file to HTML")
    parser.add_argument("input_docx", nargs="?", default="HaraldHardrada.docx", help="Input .docx file")
    parser.add_argument("output_html", nargs="?", default="harald.html", help="Output HTML file")
    args = parser.parse_args()

    docx_to_html(args.input_docx, args.output_html)

    print(f"Created {args.output_html}")



